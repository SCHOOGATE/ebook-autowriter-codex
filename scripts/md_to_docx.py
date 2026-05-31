"""
md_to_docx.py — Markdown原稿をDOCXに変換（python-docx使用）

setup.shでpython-docxがインストール済みであることが前提。
pandocは不要。

Usage:
    python scripts/md_to_docx.py output/{slug}/manuscript.md output/{slug}/manuscript.docx
"""

import re
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx がインストールされていません。")
    print("setup.sh を実行してください: pip install python-docx")
    sys.exit(1)


def parse_markdown(md_text: str) -> list:
    """Markdownをブロック単位にパース"""
    blocks = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 見出し
        if line.startswith('# '):
            blocks.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            blocks.append(('h4', line[5:].strip()))
        # 箇条書き
        elif re.match(r'^[-*] ', line):
            blocks.append(('bullet', line[2:].strip()))
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            blocks.append(('number', text))
        # 水平線（ページ区切り）
        elif re.match(r'^---+$', line.strip()):
            blocks.append(('pagebreak', ''))
        # 空行
        elif line.strip() == '':
            pass
        # 通常段落
        else:
            blocks.append(('para', line.strip()))
        i += 1

    return blocks


def apply_inline_formatting(paragraph, text: str):
    """太字・イタリックのインライン書式を適用"""
    # **太字** と *イタリック* を処理
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def convert(md_path: str, docx_path: str):
    """MarkdownファイルをDOCXに変換"""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    blocks = parse_markdown(md_text)
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10.5)
    font.name = 'Yu Mincho'

    for block_type, text in blocks:
        if block_type == 'h1':
            p = doc.add_heading(text, level=1)
            # 章の前にページ区切り（最初の見出し以外）
            if len(doc.paragraphs) > 1:
                p.paragraph_format.page_break_before = True
        elif block_type == 'h2':
            doc.add_heading(text, level=2)
        elif block_type == 'h3':
            doc.add_heading(text, level=3)
        elif block_type == 'h4':
            doc.add_heading(text, level=4)
        elif block_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text)
        elif block_type == 'number':
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text)
        elif block_type == 'pagebreak':
            doc.add_page_break()
        elif block_type == 'para':
            p = doc.add_paragraph()
            apply_inline_formatting(p, text)

    doc.save(docx_path)
    file_size = os.path.getsize(docx_path)
    print(f"DOCX生成完了: {docx_path} ({file_size:,} bytes)")


def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/md_to_docx.py input.md output.docx")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not os.path.isfile(md_path):
        print(f"ERROR: ファイルが見つかりません: {md_path}")
        sys.exit(1)

    convert(md_path, docx_path)


if __name__ == '__main__':
    main()
